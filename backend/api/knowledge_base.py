import os
import traceback
from datetime import datetime
from pathlib import Path

import pymysql
from fastapi import APIRouter, Query, HTTPException, UploadFile, File, Form

router = APIRouter(prefix="/api/knowledge-base")

UPLOAD_DIR = Path(__file__).resolve().parent.parent / "uploads"
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)

ALLOWED_EXTENSIONS = {".pdf", ".doc", ".docx", ".txt"}


def _get_conn():
    return pymysql.connect(
        host=os.getenv("MYSQL_HOST", "192.168.140.139"),
        port=int(os.getenv("MYSQL_PORT", "3306")),
        user=os.getenv("MYSQL_USER", "root"),
        password=os.getenv("MYSQL_PASSWORD", "123"),
        database=os.getenv("MYSQL_DATABASE", "medicine"),
        charset="utf8mb4",
        cursorclass=pymysql.cursors.DictCursor,
    )


def _now():
    return datetime.now().strftime("%Y-%m-%d %H:%M:%S")


# ==================== 知识库 CRUD ====================

@router.get("/list")
def list_knowledge_bases(
    page: int = Query(default=1, ge=1),
    page_size: int = Query(default=20, ge=1, le=100),
):
    conn = _get_conn()
    try:
        with conn.cursor() as cur:
            cur.execute("SELECT COUNT(*) as cnt FROM kb_base WHERE is_deleted=0")
            total = cur.fetchone()["cnt"]

            offset = (page - 1) * page_size
            cur.execute(
                "SELECT * FROM kb_base WHERE is_deleted=0 "
                "ORDER BY created_at DESC LIMIT %s OFFSET %s",
                (page_size, offset),
            )
            rows = cur.fetchall()

        for row in rows:
            row["name"] = row.pop("display_name", "")
            row["embeddingModel"] = row.pop("embedding_model", "")
            row["createTime"] = row.pop("created_at", None)
            row.pop("update_by", None)
            row.pop("create_by", None)
            row.pop("updated_at", None)
            row.pop("is_deleted", None)
            row.pop("deleted_at", None)
            row.pop("cover", None)
            row.pop("knowledge_name", None)
            row.pop("embedding_dim", None)

            with conn.cursor() as cur:
                cur.execute(
                    "SELECT COUNT(*) as cnt FROM kb_document WHERE knowledge_base_id=%s",
                    (row["id"],),
                )
                row["docCount"] = cur.fetchone()["cnt"]

        return {"code": 200, "data": rows, "total": total, "page": page, "pageSize": page_size}
    finally:
        conn.close()


@router.post("/create")
def create_knowledge_base(data: dict):
    conn = _get_conn()
    try:
        name = data.get("name", "").strip()
        if not name:
            raise HTTPException(status_code=400, detail="知识库名称不能为空")

        knowledge_name = data.get("knowledgeName", name.lower().replace(" ", "_"))
        description = data.get("description", "")
        embedding_model = data.get("embeddingModel", "text-embedding-ada-002")
        embedding_dim = data.get("embeddingDim", 1536)

        now = _now()
        with conn.cursor() as cur:
            cur.execute(
                "INSERT INTO kb_base (knowledge_name, display_name, description, embedding_model, "
                "embedding_dim, status, create_by, update_by, created_at, updated_at) "
                "VALUES (%s, %s, %s, %s, %s, 1, %s, %s, %s, %s)",
                (knowledge_name, name, description, embedding_model, embedding_dim, "admin", "admin", now, now),
            )
            conn.commit()
            return {"code": 200, "data": {"id": cur.lastrowid}}
    finally:
        conn.close()


@router.put("/{kb_id}")
def update_knowledge_base(kb_id: int, data: dict):
    conn = _get_conn()
    try:
        name = data.get("name", "").strip()
        description = data.get("description", "")
        embedding_model = data.get("embeddingModel", "")

        now = _now()
        with conn.cursor() as cur:
            cur.execute(
                "UPDATE kb_base SET display_name=%s, description=%s, embedding_model=%s, "
                "update_by=%s, updated_at=%s WHERE id=%s AND is_deleted=0",
                (name or None, description, embedding_model, "admin", now, kb_id),
            )
            conn.commit()
            return {"code": 200}
    finally:
        conn.close()


@router.delete("/{kb_id}")
def delete_knowledge_base(kb_id: int):
    conn = _get_conn()
    try:
        now = _now()
        with conn.cursor() as cur:
            cur.execute(
                "UPDATE kb_base SET is_deleted=1, deleted_at=%s WHERE id=%s",
                (now, kb_id),
            )
            conn.commit()
            return {"code": 200}
    finally:
        conn.close()


# ==================== 文件解析 ====================

def _parse_pdf(file_path: str) -> str:
    try:
        from pypdf import PdfReader
        reader = PdfReader(file_path)
        parts = [(page.extract_text() or "").strip() for page in reader.pages]
        return "\n\n".join(part for part in parts if part)
    except ImportError:
        raise HTTPException(status_code=500, detail="pypdf 未安装，无法解析 PDF 文件")


def _parse_docx(file_path: str) -> str:
    try:
        from docx import Document
        doc = Document(file_path)
        parts = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
        return "\n\n".join(parts)
    except ImportError:
        raise HTTPException(status_code=500, detail="python-docx 未安装，无法解析 Word 文件")


def _parse_doc(file_path: str) -> str:
    try:
        from docx import Document
        doc = Document(file_path)
        parts = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
        return "\n\n".join(parts)
    except Exception:
        try:
            import subprocess
            result = subprocess.run(
                ["antiword", file_path],
                capture_output=True, text=True, timeout=30,
            )
            if result.returncode == 0 and result.stdout.strip():
                return result.stdout.strip()
        except Exception:
            pass
        raise HTTPException(
            status_code=400,
            detail="无法解析 .doc 文件，请将其转换为 .docx 格式后重试",
        )


def _parse_txt(file_path: str) -> str:
    encodings = ["utf-8", "gbk", "gb2312", "latin-1"]
    for enc in encodings:
        try:
            with open(file_path, "r", encoding=enc) as f:
                return f.read().strip()
        except UnicodeDecodeError:
            continue
    raise HTTPException(status_code=400, detail="无法识别文件编码，请使用 UTF-8 编码")


def _parse_file(file_path: str, file_type: str) -> str:
    ext = file_type.lower()
    if ext == ".pdf":
        return _parse_pdf(file_path)
    elif ext == ".docx":
        return _parse_docx(file_path)
    elif ext == ".doc":
        return _parse_doc(file_path)
    elif ext == ".txt":
        return _parse_txt(file_path)
    else:
        raise HTTPException(status_code=400, detail=f"不支持的文件类型: {ext}")


# ==================== 文档管理 ====================

@router.get("/{kb_id}/documents")
def list_documents(
    kb_id: int,
    page: int = Query(default=1, ge=1),
    page_size: int = Query(default=20, ge=1, le=100),
):
    conn = _get_conn()
    try:
        with conn.cursor() as cur:
            cur.execute(
                "SELECT COUNT(*) as cnt FROM kb_document WHERE knowledge_base_id=%s",
                (kb_id,),
            )
            total = cur.fetchone()["cnt"]

            offset = (page - 1) * page_size
            cur.execute(
                "SELECT * FROM kb_document WHERE knowledge_base_id=%s "
                "ORDER BY created_at DESC LIMIT %s OFFSET %s",
                (kb_id, page_size, offset),
            )
            rows = cur.fetchall()

        for row in rows:
            row["fileName"] = row.pop("file_name", "")
            row["fileSize"] = row.pop("file_size", 0)
            row["fileType"] = row.pop("file_type", "")
            row["status"] = row.pop("stage", "PENDING")
            row["createTime"] = row.pop("created_at", None)
            row.pop("file_url", None)
            row.pop("chunk_mode", None)
            row.pop("chunk_size", None)
            row.pop("chunk_overlap", None)
            row.pop("last_error", None)
            row.pop("create_by", None)
            row.pop("update_by", None)
            row.pop("updated_at", None)

            with conn.cursor() as cur:
                cur.execute(
                    "SELECT COUNT(*) as cnt FROM kb_document_chunk WHERE document_id=%s",
                    (row["id"],),
                )
                row["chunkCount"] = cur.fetchone()["cnt"]

        return {"code": 200, "data": rows, "total": total, "page": page, "pageSize": page_size}
    finally:
        conn.close()


@router.post("/upload")
async def upload_document(
    file: UploadFile = File(...),
    knowledgeBaseId: int = Form(...),
):
    if not file.filename:
        raise HTTPException(status_code=400, detail="未选择文件")

    ext = Path(file.filename).suffix.lower()
    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=400,
            detail=f"不支持的文件类型 {ext}，仅支持: {', '.join(ALLOWED_EXTENSIONS)}",
        )

    conn = _get_conn()
    try:
        with conn.cursor() as cur:
            cur.execute(
                "SELECT id FROM kb_base WHERE id=%s AND is_deleted=0",
                (knowledgeBaseId,),
            )
            if not cur.fetchone():
                raise HTTPException(status_code=404, detail="知识库不存在")
    finally:
        conn.close()

    now = _now()
    timestamp = datetime.now().strftime("%Y%m%d%H%M%S%f")
    safe_filename = f"{timestamp}_{file.filename}"
    dest_path = UPLOAD_DIR / safe_filename
    file_size = 0

    try:
        content = await file.read()
        file_size = len(content)
        with open(dest_path, "wb") as f:
            f.write(content)
    except Exception as e:
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"文件保存失败: {str(e)}")

    file_url = f"/uploads/{safe_filename}"

    try:
        parsed_text = _parse_file(str(dest_path), ext)
    except HTTPException:
        raise
    except Exception as e:
        traceback.print_exc()
        if dest_path.exists():
            dest_path.unlink()
        raise HTTPException(status_code=500, detail=f"文件解析失败: {str(e)}")

    conn = _get_conn()
    try:
        with conn.cursor() as cur:
            cur.execute(
                "INSERT INTO kb_document (knowledge_base_id, file_name, file_url, file_type, "
                "file_size, stage, create_by, update_by, created_at, updated_at) "
                "VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s)",
                (knowledgeBaseId, file.filename, file_url, ext, file_size,
                 "COMPLETED", "admin", "admin", now, now),
            )
            doc_id = cur.lastrowid
            conn.commit()

        _create_chunks(conn, knowledgeBaseId, doc_id, parsed_text)

        return {"code": 200, "data": {"id": doc_id, "fileName": file.filename}}
    finally:
        conn.close()


def _create_chunks(conn, kb_id: int, doc_id: int, text: str, chunk_size: int = 500):
    if not text or not text.strip():
        return

    chunks = _split_text(text, chunk_size)
    now = _now()

    with conn.cursor() as cur:
        for idx, chunk_text in enumerate(chunks):
            if not chunk_text.strip():
                continue
            cur.execute(
                "INSERT INTO kb_document_chunk (knowledge_base_id, document_id, chunk_index, "
                "content, char_count, status, stage, created_at, updated_at) "
                "VALUES (%s, %s, %s, %s, %s, 0, 'COMPLETED', %s, %s)",
                (kb_id, doc_id, idx, chunk_text, len(chunk_text), now, now),
            )
        conn.commit()

    _vectorize_chunks(conn, kb_id, doc_id)


def _split_text(text: str, chunk_size: int = 500) -> list[str]:
    if len(text) <= chunk_size:
        return [text]

    chunks = []
    paragraphs = text.split("\n\n")
    current = ""
    for para in paragraphs:
        if len(current) + len(para) <= chunk_size:
            current = (current + "\n\n" + para).strip()
        else:
            if current:
                chunks.append(current)
            current = para
    if current:
        chunks.append(current)

    if not chunks:
        chunks = [text[i:i + chunk_size] for i in range(0, len(text), chunk_size)]

    return chunks


def _vectorize_chunks(conn, kb_id: int, doc_id: int):
    try:
        from rag.rag import _get_embedding, _get_milvus_client
        client = _get_milvus_client()
        collection = os.getenv("MILVUS_COLLECTION", "medicine_knowledge")

        with conn.cursor() as cur:
            cur.execute(
                "SELECT id, content, chunk_index FROM kb_document_chunk "
                "WHERE document_id=%s ORDER BY chunk_index",
                (doc_id,),
            )
            chunks = cur.fetchall()

        for chunk in chunks:
            try:
                embedding = _get_embedding(chunk["content"])
                data = {
                    "id": chunk["id"],
                    "document_id": doc_id,
                    "knowledge_base_id": kb_id,
                    "chunk_index": chunk["chunk_index"],
                    "content": chunk["content"],
                    "embedding": embedding,
                }
                client.insert(collection_name=collection, data=[data])
                with conn.cursor() as cur:
                    cur.execute(
                        "UPDATE kb_document_chunk SET vector_id=%s WHERE id=%s",
                        (str(chunk["id"]), chunk["id"]),
                    )
                    conn.commit()
            except Exception as e:
                print(f"向量化 chunk {chunk['id']} 失败: {e}")
    except ImportError:
        print("Milvus 或 embedding 模块未就绪，跳过向量化")
    except Exception as e:
        print(f"向量化过程出错: {e}")


@router.delete("/{kb_id}/documents/{doc_id}")
def delete_document(kb_id: int, doc_id: int):
    conn = _get_conn()
    try:
        with conn.cursor() as cur:
            cur.execute(
                "SELECT file_url FROM kb_document WHERE id=%s AND knowledge_base_id=%s",
                (doc_id, kb_id),
            )
            row = cur.fetchone()
            if not row:
                raise HTTPException(status_code=404, detail="文档不存在")

            file_url = row.get("file_url", "")
            if file_url:
                file_path = UPLOAD_DIR / Path(file_url).name
                if file_path.exists():
                    file_path.unlink()

            cur.execute("DELETE FROM kb_document_chunk WHERE document_id=%s", (doc_id,))
            cur.execute("DELETE FROM kb_document WHERE id=%s", (doc_id,))
            conn.commit()

        try:
            from rag.rag import _get_milvus_client
            client = _get_milvus_client()
            collection = os.getenv("MILVUS_COLLECTION", "medicine_knowledge")
            client.delete(collection_name=collection, filter=f'document_id == {doc_id}')
        except Exception as e:
            print(f"Milvus 删除向量失败: {e}")

        return {"code": 200}
    finally:
        conn.close()


@router.post("/{kb_id}/documents/{doc_id}/reparse")
def reparse_document(kb_id: int, doc_id: int):
    conn = _get_conn()
    try:
        with conn.cursor() as cur:
            cur.execute(
                "SELECT * FROM kb_document WHERE id=%s AND knowledge_base_id=%s",
                (doc_id, kb_id),
            )
            doc = cur.fetchone()
            if not doc:
                raise HTTPException(status_code=404, detail="文档不存在")

            file_url = doc.get("file_url", "")
            file_path = UPLOAD_DIR / Path(file_url).name if file_url else None

            if not file_path or not file_path.exists():
                raise HTTPException(status_code=404, detail="文件不存在，请重新上传")

            cur.execute("DELETE FROM kb_document_chunk WHERE document_id=%s", (doc_id,))
            conn.commit()

            try:
                from rag.rag import _get_milvus_client
                client = _get_milvus_client()
                collection = os.getenv("MILVUS_COLLECTION", "medicine_knowledge")
                client.delete(collection_name=collection, filter=f'document_id == {doc_id}')
            except Exception as e:
                print(f"Milvus 删除向量失败: {e}")

        ext = doc.get("file_type", "")
        parsed_text = _parse_file(str(file_path), ext)

        _create_chunks(conn, kb_id, doc_id, parsed_text)

        now = _now()
        with conn.cursor() as cur:
            cur.execute(
                "UPDATE kb_document SET stage='COMPLETED', updated_at=%s WHERE id=%s",
                (now, doc_id),
            )
            conn.commit()

        return {"code": 200}
    finally:
        conn.close()